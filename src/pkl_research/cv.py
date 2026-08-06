"""Analisa CV: ekstraksi teks, skor arah, dan checklist ATS (murni)."""

from __future__ import annotations

import json
import re
from pathlib import Path

_SKILLS: dict[str, dict[str, int]] = {
    "software": {
        "python": 2, "typescript": 2, "javascript": 2, "rust": 2, "golang": 2,
        "php": 2, "java": 2, "cpp": 2, "git": 1, "docker": 1, "oop": 2,
        "testing": 2, "sql": 2, "api": 1, "linux": 1, "system": 1,
    },
    "ai": {
        "python": 3, "machine learning": 3, "deep learning": 3, "llm": 3,
        "tensorflow": 3, "pytorch": 3, "scikit": 2, "keras": 2, "nlp": 2,
        "computer vision": 2, "openai": 2, "gemini": 2, "chromadb": 2,
        "vector database": 2, "prompt": 2, "claude": 2, "opencode": 2,
        "chatbot": 2, "agent": 1, "data": 1, "rag": 2, "multi-agent": 2,
    },
    "fullstack": {
        "next.js": 3, "react": 3, "vue": 2, "node.js": 3, "express": 2,
        "django": 2, "laravel": 2, "typescript": 2, "javascript": 2,
        "frontend": 3, "backend": 3, "api": 2, "database": 2, "mysql": 2,
        "postgres": 2, "tailwind": 1, "docker": 1, "rest": 1,
    },
    "game": {
        "unity": 3, "unreal": 3, "godot": 3, "gdscript": 3, "c#": 2,
        "game": 2, "gamedev": 2, "blender": 2, "3d": 1, "hackathon": 1,
        "game design": 2,
    },
}

_ALIASES = {"c++": "cpp", "c": "c"}

ROLE_LABEL = {
    "software": "Software Developer",
    "ai": "AI Developer",
    "fullstack": "Fullstack Developer",
    "game": "Game Developer",
}

_ROLE_THRESHOLD = {"software": 26, "ai": 18, "fullstack": 22, "game": 16}


def extract_text(path: str | Path) -> str:
    """Ekstrak teks dari PDF/DOCX/TXT."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        parts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    return path.read_text(encoding="utf-8", errors="replace")


def _norm(text: str) -> str:
    return text.lower()


def _kw_found(text: str, keyword: str) -> bool:
    keyword = _ALIASES.get(keyword, keyword)
    if " " in keyword:
        return keyword in text
    return re.search(rf"\b{re.escape(keyword)}\b", text) is not None


def _found_weighted(text: str, keywords: dict[str, int]) -> tuple[int, int, list[str]]:
    found_weight = 0
    found_words: list[str] = []
    for kw, weight in keywords.items():
        if _kw_found(text, kw):
            found_weight += weight
            found_words.append(kw)
    return found_weight, sum(keywords.values()), found_words


def _pct(found: int, threshold: int) -> int:
    return min(100, round(found / threshold * 100))


def analyze_cv(text: str) -> dict:
    """Analisa teks CV → skor 4 arah + skill + checklist ATS."""
    norm = _norm(text)
    summary = norm[:700]

    scores: dict[str, int] = {}
    skills: dict[str, list[str]] = {}
    for direction, keywords in _SKILLS.items():
        found_w, _, found_words = _found_weighted(norm, keywords)
        obj_w, _, _ = _found_weighted(summary, keywords)
        skill_pct = _pct(found_w, _ROLE_THRESHOLD[direction])
        obj_pct = _pct(obj_w, 6)
        scores[direction] = round(0.65 * skill_pct + 0.35 * obj_pct)
        skills[direction] = found_words

    strengths, gaps = _strengths_gaps(scores, skills)
    ats = ats_checks(text)
    return {
        "scores": scores,
        "skills": skills,
        "strengths": strengths,
        "gaps": gaps,
        "ats": ats,
    }


def _strengths_gaps(scores: dict[str, int], skills: dict[str, list[str]]) -> tuple[list[str], list[str]]:
    top = sorted(scores.items(), key=lambda kv: -kv[1])[:2]
    strengths = [
        f"Paling cocok ke arah **{ROLE_LABEL[d]}** (skor {s}/100) — "
        f"didukung skill: {', '.join(skills[d][:6]) or '-'}"
        for d, s in top
    ]
    low = [d for d, s in scores.items() if s < 50]
    gaps = []
    for d in low:
        found = set(skills[d])
        missing = [kw for kw in _SKILLS[d] if kw not in found][:6]
        if missing:
            gaps.append(
                f"Untuk **{ROLE_LABEL[d]}** (skor {scores[d]}/100) kurang signal: "
                f"tambah {', '.join(missing)}"
            )
    return strengths, gaps


def ats_checks(text: str) -> list[dict]:
    """Checklist ATS-friendliness. Return list of {check, ok, note}."""
    norm = _norm(text)
    checks: list[dict] = []
    checks.append({
        "check": "Kontak lengkap (email + telepon)",
        "ok": bool(re.search(r"[\w.+-]+@[\w.-]+", text)) and bool(re.search(r"(\+?\d[\d\s-]{8,})", text)),
        "note": "",
    })
    checks.append({
        "check": "Section jelas (Education / Skills / Experience)",
        "ok": all(s in norm for s in ("education", "skill")) and any(
            s in norm for s in ("projects", "experience", "pengalaman")
        ),
        "note": "",
    })
    checks.append({
        "check": "Kata kunci teknis relevan (ATS keyword)",
        "ok": sum(1 for _d, kws in _SKILLS.items() for kw in kws if _kw_found(norm, kw)) >= 10,
        "note": ">=10 keyword teknis terdeteksi",
    })
    numbers = re.findall(r"\d[\d,.]*", text)
    checks.append({
        "check": "Ada pencapaian terukur (angka)",
        "ok": len(numbers) >= 4,
        "note": f"{len(numbers)} angka terdeteksi",
    })
    words = len(re.findall(r"\S+", text))
    checks.append({
        "check": "Panjang ideal (1-2 halaman, ~200-800 kata)",
        "ok": 200 <= words <= 800,
        "note": f"{words} kata",
    })
    return checks


def save_analysis(analysis: dict, path: str | Path) -> None:
    Path(path).write_text(json.dumps(analysis, ensure_ascii=False, indent=2), encoding="utf-8")


def load_analysis(path: str | Path) -> dict | None:
    p = Path(path)
    if not p.exists():
        return None
    try:
        return json.loads(p.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return None


def fit_for_roles(analysis: dict, role_fit: list[str]) -> float:
    """Skor fit perusahaan (mean skor CV untuk role_fit-nya)."""
    scores = analysis.get("scores", {})
    if not role_fit:
        return 0.0
    relevant = [scores.get(r, 0) for r in role_fit if r in scores]
    if not relevant:
        return 0.0
    return round(sum(relevant) / len(relevant), 1)
